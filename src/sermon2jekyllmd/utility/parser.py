from docx import Document
from . import md_gen
import re
from .replacer import Replacer

# pip install docx-parser

def replace_incorrect_chars(text):
    """
    Replace known incorrect characters with correct characters.
    """
    replacer = Replacer()
    return replacer.replace_characters(text)

def read_docx(file_path):
    """
    Read and extract text from a .docx file.

    Parameters:
        file_path (str): The path to the .docx file.

    Returns:
        str: The extracted text from the .docx file.
    """
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


#
#
# # Example usage:
# file_path = "./docs/74-75.docx"  # Replace with the path to your .docx file
# extracted_text = read_docx(file_path)
# print(extracted_text)
#

def read_table(doc):
    """
    Read and extract data from tables in a .docx file.

    Parameters:
        doc (Document): The Document object representing the .docx file.

    Returns:
        list: A list of lists, where each inner list represents a row in the table.
    """
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(md_gen.gen_html_underline_from_cell(cell))
            table_data.append(row_data)
    return table_data


def read_doc_table(file_path):
    doc = Document(file_path)
    table_data = read_table(doc)
    return table_data


# Example usage:
# file_path = "./docs/76-77.docx"  # Replace with the path to your .docx file
# doc = Document(file_path)
# table_data = read_table(doc)

def parse_table_length_12(table_data):
    weekday_titles = []
    weekday_contents = []
    header123 = table_data[0]
    title = header123[0].strip()
    #print("Titile is" + title)
    for i in [3, 7, 10]:
        weekday_titles.append(header123[i].strip())
    content123 = table_data[1]
    #print(content123)
    for i in [3, 7, 10]:
        #weekday_contents.append(content123[i].strip())
        weekday_contents.append(content123[i])
    header456 = table_data[2]
    for i in [4, 7, 10]:
        weekday_titles.append(header456[i].strip())
    content456 = table_data[3]
    for i in [4, 7, 10]:
        #weekday_contents.append(content456[i].strip())
        weekday_contents.append(content456[i])
    return title, weekday_titles, weekday_contents

def parse_table_length_11(table_data):
    weekday_titles = []
    weekday_contents = []
    header123 = table_data[0]
    title = header123[0].strip()
    #print("Titile is" + title)
    for i in [3, 7, 10]:
        weekday_titles.append(header123[i].strip())
    content123 = table_data[1]
    #print(content123)
    for i in [3, 7, 10]:
        #weekday_contents.append(content123[i].strip())
        weekday_contents.append(content123[i])
    header456 = table_data[2]
    for i in [4, 7, 10]:
        weekday_titles.append(header456[i].strip())
    content456 = table_data[3]
    for i in [4, 7, 10]:
        #weekday_contents.append(content456[i].strip())
        weekday_contents.append(content456[i])
    return title, weekday_titles, weekday_contents



def parse_table_length_7(table_data):
    weekday_titles = []
    weekday_contents = []
    header123 = table_data[0]
    title = header123[0].strip()
    #print("Titile is" + title)
    for i in [2, 4, 6]:
        weekday_titles.append(header123[i].strip())
    content123 = table_data[1]
    #print(content123)
    for i in [1, 3, 5]:
        #weekday_contents.append(content123[i].strip())
        weekday_contents.append(content123[i])
    header456 = table_data[2]
    #print(header456)
    if len(header456) == 5:
        for i in [2, 3, 4]:
            weekday_titles.append(header456[i].strip())
    else:
        for i in [2, 4, 6]:
            weekday_titles.append(header456[i].strip())
    content456 = table_data[3]
    #print(content456)
    if len(header456) == 5:
        for i in [2, 3, 4]:
            #weekday_contents.append(content456[i].strip())
            weekday_contents.append(content456[i])
    else:
        for i in [1, 3, 5]:
            #weekday_contents.append(content456[i].strip())
            weekday_contents.append(content456[i])
    return title, weekday_titles, weekday_contents


def parse_table_length_4(table_data):
    weekday_titles = []
    weekday_contents = []
    title = ''
    header123 = table_data[0]
    title = header123.pop(0)
    for i in header123:
        header_temp = i.split()
        weekday_titles.append(header_temp[-1].strip())
    content123 = table_data[1]
    content123.pop(0)
    for i in content123:
        weekday_contents.append(i)
    header456 = table_data[2]
    header456.pop(0)
    for i in header456:
        header_temp = i.split()
        weekday_titles.append(header_temp[-1].strip())
    content456 = table_data[3]
    content456.pop(0)
    for i in content456:
        weekday_contents.append(i)
    return title, weekday_titles, weekday_contents

def get_bold_list(para):
    bold_list = []
    for run in para.runs:
        if run.bold:
            bold_list.append(run.text)
    return bold_list

def get_underline_list(para):
    underline_list = []
    for run in para.runs:
        if run.underline:
            underline_list.append(run.text)
    print(underline_list)
    #print(''.join(underline_list))
    return underline_list
